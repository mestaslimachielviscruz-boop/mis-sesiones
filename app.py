import io
import json
import os
import streamlit as st
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor
from google import genai
from google.genai import types

st.set_page_config(page_title="Generador CNEB", page_icon="📚", layout="centered")

SYSTEM_PROMPT = """
Eres un especialista en pedagogía y currículo del Ministerio de Educación del Perú (CNEB).
Diseña sesiones de aprendizaje con enfoque de alta demanda cognitiva.
DEBES RESPONDER ÚNICAMENTE EN FORMATO JSON ESTRUCTURADO CON ESTA ESTRUCTURA EXACTA:
{
  "datos_generales": { "titulo_sesion": "...", "area": "...", "grado": "...", "duracion": "..." },
  "propósito_aprendizaje": { "competencia": "...", "capacidades": ["..."], "desempeno_precisado": "...", "evidencia": "..." },
  "secuencia_didactica": {
    "inicio": "...",
    "desarrollo": "...",
    "cierre": {
      "actividad_cierre": "...",
      "evaluacion": {
        "retroalimentacion_formativa": { "preguntas_clarificacion": "...", "valoracion_logros": "...", "sugerencias_mejora": "..." },
        "rubrica_analitica_cneb": [ { "competencia_capacidad": "...", "criterio_evaluacion": "...", "inicio_c": "...", "proceso_b": "...", "logrado_a": "...", "destacado_ad": "..." } ]
      }
    }
  }
}
"""

def aplicar_estilo_tabla_minedu(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i == 0:
                tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="003366"/>'))
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(10)
                        r.font.color.rgb = RGBColor(255, 255, 255)
            else:
                bg = "F4F6F9" if i % 2 == 0 else "FFFFFF"
                tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}"/>'))
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9.5)

def generar_documento_bytes(data: dict) -> bytes:
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    datos_gen = data.get("datos_generales", {})
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_titulo.add_run(f"SESIÓN DE APRENDIZAJE\n“{datos_gen.get('titulo_sesion', '').upper()}”")
    run_t.bold = True
    run_t.font.size = Pt(14)
    run_t.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_heading("1. DATOS GENERALES", level=2)
    t_datos = doc.add_table(rows=2, cols=4)
    t_datos.cell(0, 0).paragraphs[0].add_run("Área:").bold = True
    t_datos.cell(0, 1).text = datos_gen.get("area", "-")
    t_datos.cell(0, 2).paragraphs[0].add_run("Grado:").bold = True
    t_datos.cell(0, 3).text = datos_gen.get("grado", "-")
    t_datos.cell(1, 0).paragraphs[0].add_run("Duración:").bold = True
    t_datos.cell(1, 1).text = datos_gen.get("duracion", "-")
    t_datos.cell(1, 2).paragraphs[0].add_run("Enfoque:").bold = True
    t_datos.cell(1, 3).text = "CNEB"
    aplicar_estilo_tabla_minedu(t_datos)

    doc.add_heading("2. PROPÓSITOS DE APRENDIZAJE", level=2)
    proposito = data.get("propósito_aprendizaje", {})
    t_prop = doc.add_table(rows=1, cols=3)
    t_prop.cell(0, 0).text = "Competencia y Capacidades"
    t_prop.cell(0, 1).text = "Desempeño Precisado"
    t_prop.cell(0, 2).text = "Evidencia"
    r_prop = t_prop.add_row()
    r_prop.cells[0].text = f"{proposito.get('competencia', '-')}\n" + "\n".join(proposito.get('capacidades', []))
    r_prop.cells[1].text = proposito.get('desempeno_precisado', '-')
    r_prop.cells[2].text = proposito.get('evidencia', '-')
    aplicar_estilo_tabla_minedu(t_prop)

    doc.add_heading("3. SECUENCIA DIDÁCTICA", level=2)
    secuencia = data.get("secuencia_didactica", {})
    t_sec = doc.add_table(rows=1, cols=2)
    t_sec.cell(0, 0).text = "Momento"
    t_sec.cell(0, 1).text = "Estrategias"
    cierre_data = secuencia.get("cierre", {})
    for mom, text in [("INICIO", secuencia.get("inicio", "-")), ("DESARROLLO", secuencia.get("desarrollo", "-")), ("CIERRE", cierre_data.get("actividad_cierre", "-") if isinstance(cierre_data, dict) else str(cierre_data))]:
        row = t_sec.add_row()
        row.cells[0].paragraphs[0].add_run(mom).bold = True
        row.cells[1].text = text
    aplicar_estilo_tabla_minedu(t_sec)

    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()

st.title("📚 Generador de Sesiones CNEB")

api_key = os.environ.get("GEMINI_API_KEY")
if not api_key:
    api_key = st.text_input("🔑 Ingrese su clave de API de Gemini:", type="password")

area = st.selectbox("📌 Área Curricular", ["Comunicación", "Desarrollo Personal, Ciudadanía y Cívica (DPCC)", "Ciencias Sociales", "Ciencia y Tecnología", "Matemática"])
grado = st.selectbox("🎓 Grado", ["1° de Secundaria", "2° de Secundaria", "3° de Secundaria", "4° de Secundaria", "5° de Secundaria"])
tema = st.text_area("✍️ Tema de la sesión:")
duracion = st.text_input("⏱️ Duración:", value="90 minutos")

if st.button("🚀 Generar Sesión en Word", type="primary"):
    if not api_key or not tema.strip():
        st.warning("Ingrese su clave y el tema.")
    else:
        with st.spinner("Creando sesión..."):
            try:
                client = genai.Client(api_key=api_key)
                resp = client.models.generate_content(
                    model="gemini-2.5-pro",
                    contents=f"Área: {area}, Grado: {grado}, Tema: {tema}, Duración: {duracion}",
                    config=types.GenerateContentConfig(system_instruction=SYSTEM_PROMPT, response_mime_type="application/json", temperature=0.2)
                )
                docx_bytes = generar_documento_bytes(json.loads(resp.text))
                st.success("¡Listo!")
                st.download_button("📥 Descargar Word (.docx)", data=docx_bytes, file_name=f"Sesion_{area}_{grado}.docx")
            except Exception as e:
                st.error(f"Error: {e}")
